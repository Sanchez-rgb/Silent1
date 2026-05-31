"""生成测试用的 DOCX 文件"""
try:
    from docx import Document
    has_docx = True
except ImportError:
    has_docx = False
    print("警告: python-docx 未安装，将使用模拟测试文件")

import os
import zipfile
from datetime import datetime

def create_test_docx(filename, content):
    """创建一个简单的 DOCX 文件"""
    if not has_docx:
        print(f"无法创建真实的 DOCX 文件: {filename}")
        print("正在创建空的占位文件用于测试...")
        os.makedirs('test_data', exist_ok=True)
        filepath = os.path.join('test_data', filename)
        # 创建一个简单的文本文件作为占位符
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"测试文件: {filename}\n")
            f.write(f"创建时间: {datetime.now()}\n")
            f.write(f"内容: {content}\n")
        print(f'创建占位文件: {filepath}')
        return filepath
    
    doc = Document()
    doc.add_heading('Word 转 PDF 测试文件', 0)
    doc.add_paragraph('这是一个用于测试 Word 转 PDF 服务的文档。')
    
    doc.add_heading('测试内容', level=1)
    doc.add_paragraph(content)
    
    doc.add_heading('功能列表', level=2)
    doc.add_paragraph('✅ 支持单个文件转换', style='List Bullet')
    doc.add_paragraph('✅ 支持批量转换', style='List Bullet')
    doc.add_paragraph('✅ 自动 ZIP 打包', style='List Bullet')
    
    doc.add_paragraph(f'测试时间: {datetime.now()}')
    
    os.makedirs('test_data', exist_ok=True)
    filepath = os.path.join('test_data', filename)
    doc.save(filepath)
    print(f'创建测试文件: {filepath}')
    return filepath

if __name__ == '__main__':
    print('正在生成测试文件...')
    try:
        files = []
        files.append(create_test_docx('test_file_1.docx', '第一个测试文档的内容，用于验证转换服务是否正常工作。'))
        files.append(create_test_docx('test_file_2.docx', '第二个测试文档，包含一些示例内容，帮助测试批量转换功能。'))
        files.append(create_test_docx('test_file_3.docx', '第三个测试文档，确保我们的服务能够处理多个文件的转换请求。'))
        print('\n测试文件生成完成！')
        print(f'生成的文件:')
        for f in files:
            if os.path.exists(f):
                print(f'  ✅ {os.path.abspath(f)} ({os.path.getsize(f)} bytes)')
    except Exception as e:
        print(f'错误: {e}')
        import traceback
        traceback.print_exc()
